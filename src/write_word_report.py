"""Generate a detailed Word report from pipeline outputs."""

from __future__ import annotations

import json
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from data_loader import FIG_DIR, OUT_DIR, REPORT_DIR
from report_stats import archetype_definitions, archetype_lines, feature_description, pre_start_validation, split_summary


def _load_json(name: str) -> dict[str, Any]:
    return json.loads((OUT_DIR / name).read_text())


def _load_csv(name: str) -> pd.DataFrame:
    return pd.read_csv(OUT_DIR / name)


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def _add_body(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _add_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Light Grid Accent 1"
    for j, header in enumerate(headers):
        table.rows[0].cells[j].text = str(header)
    for i, row in enumerate(rows, start=1):
        for j, value in enumerate(row):
            table.rows[i].cells[j].text = str(value)
    doc.add_paragraph("")


def _add_figure(doc: Document, path: Path, caption: str, width: float = 5.5) -> None:
    if path.exists():
        doc.add_picture(str(path), width=Inches(width))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.runs[0].italic = True
        doc.add_paragraph("")


def write_word_report() -> Path:
    REPORT_DIR.mkdir(parents=True, exist_ok=True)
    output_path = REPORT_DIR / "Studor_PathAI_Report.docx"

    weights_df = _load_csv("engagement_weight_rationale.csv")
    score_bands = _load_csv("engagement_score_band_risk.csv")
    feature_rationale = _load_csv("feature_rationale_week6.csv")
    consistency = _load_csv("data_consistency_audit.csv")
    leakage = _load_csv("leakage_audit.csv")
    overview = _load_csv("data_cleaning_overview.csv")
    model_comparison = _load_csv("model_comparison_cv.csv")
    threshold_table = _load_csv("risk_threshold_analysis.csv")
    intervention_tiers = _load_csv("risk_intervention_tiers.csv")
    behavioral_drivers = _load_csv("risk_behavioral_feature_drivers.csv")
    rec_metrics = _load_json("recommendation_metrics.json")
    urgent = _load_json("risk_metrics.json")[0]
    watchlist = _load_json("risk_watchlist_metrics.json")[0]
    pre_start = pre_start_validation()
    split = split_summary()
    assess_row = feature_rationale.loc[feature_rationale["feature"] == "assessment_submitted_ratio"].iloc[0]
    engagement_row = feature_rationale.loc[feature_rationale["feature"] == "engagement_score"].iloc[0]
    neg_date_count = int(
        consistency.loc[consistency["check"].str.contains("Negative VLE", na=False), "issue_count"].iloc[0]
    )

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    title = doc.add_heading("Studor PathAI: Technical Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Open University Learning Analytics Dataset (OULAD) — Behavioral Scoring, "
        "Week 6 Disengagement Prediction, and Course Recommendation"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("")

    _add_heading(doc, "Executive Summary", 1)
    _add_body(
        doc,
        "This report documents the end-to-end PathAI prototype built on one semester of OULAD behavioral data. "
        "The pipeline converts raw clickstream, assessment, and demographic records into three product-facing assets: "
        "(1) a weekly 0–100 engagement score; (2) a Week 6 disengagement risk model; "
        "and (3) a next-module recommendation engine."
    )
    _add_heading(doc, "Task 1: Behavioral Engagement Scoring", 2)
    _add_body(
        doc,
        "The score is built from six interpretable feature buckets: activity volume, consistency and rhythm, "
        "recency gaps, resource diversity and intent, course-pace alignment, and assessment behavior. These signals "
        "support seven archetypes: steady engager, early dropout, late recoverer, burst engager, perfectionist "
        "procrastinator, compliance engager, and opportunistic engager. "
        f"The strongest validation result is the Week 6 score gradient: {score_bands.iloc[0]['observed_withdraw_fail_rate']:.1%} "
        f"withdraw/fail in the 0-20 score band versus {score_bands.iloc[-1]['observed_withdraw_fail_rate']:.1%} in the 80-100 band."
    )
    _add_heading(doc, "Task 2: Week 6 Disengagement Model", 2)
    _add_body(
        doc,
        "The final model is XGBoost, trained only on features available by Week 6. Its output is framed as three "
        "intervention tiers rather than a single urgent alert, reducing alert fatigue while preserving actionability. "
        f"The headline metrics are ROC-AUC {urgent['roc_auc']:.3f}, high-touch precision {urgent['precision']:.3f}, "
        f"top-60% light-touch F1 {watchlist['f1']:.3f}, and top-60% light-touch recall {watchlist['recall']:.3f}. "
        "Model selection and alerting are separated deliberately: XGBoost ranks risk, while intervention thresholds "
        "reflect advisor capacity and alert-fatigue risk."
    )
    _add_heading(doc, "Task 3: Next-Module Recommendation", 2)
    _add_body(
        doc,
        "The recommender compares collaborative filtering against a shared student-course feature-space content model. "
        "The content model uses cosine similarity plus a Wilson lower-bound pass-rate prior, which keeps cold-start "
        "recommendations conservative instead of chasing noisy pass rates. "
        f"Content hit@3 is {rec_metrics['content_hit_rate_at_3']:.3f}, collaborative hit@3 is "
        f"{rec_metrics['cf_hit_rate_at_3']:.3f}, and the cold-start modules are {', '.join(rec_metrics['cold_start_strategy'])}."
    )
    _add_body(
        doc,
        f"The dataset contains {overview.loc[overview['table'] == 'student_info', 'rows'].iloc[0]:,.0f} enrolments "
        f"across 7 modules. Overall risk rate is {split['risk_rate']:.1%}. The core product choice is to translate "
        "statistical outputs into advisor-friendly workflows rather than exposing raw model scores alone."
    )

    _add_heading(doc, "1. Data Foundation and Cleaning Decisions", 1)
    _add_heading(doc, "1.1 Enrolment Grain", 2)
    _add_body(
        doc,
        "The enrolment key is (id_student, code_module, code_presentation). Student ID alone is not unique because "
        "learners can take multiple modules simultaneously."
    )

    _add_heading(doc, "1.2 Table Inventory", 2)
    overview_rows = [
        [row["table"], f"{int(row['rows']):,}", int(row["columns"]), int(row["missing_cells"])]
        for _, row in overview.iterrows()
    ]
    _add_table(doc, ["Table", "Rows", "Columns", "Missing cells"], overview_rows)

    _add_heading(doc, "1.3 Consistency Checks", 2)
    consistency_rows = [[row["check"], int(row["issue_count"]), row["policy"]] for _, row in consistency.iterrows()]
    _add_table(doc, ["Check", "Issue count", "Policy"], consistency_rows)
    _add_body(
        doc,
        f"{neg_date_count:,} VLE events have negative dates — valid pre-presentation activity before module day 0. "
        "These feed proactivity features across all three tasks. In-week aggregates still use day 0+ only."
    )

    _add_heading(doc, "1.4 Leakage Controls", 2)
    leakage_rows = [[row["candidate_feature"], row["available_by_week6"], row["used"], row["reason"]] for _, row in leakage.iterrows()]
    _add_table(doc, ["Candidate feature", "Available Week 6?", "Used?", "Reason"], leakage_rows)
    _add_figure(doc, FIG_DIR / "outcome_eda.png", "Figure 1. Outcome distribution and withdraw/fail rate by module.")
    _add_figure(doc, FIG_DIR / "data_missingness.png", "Figure 2. Top missingness patterns.")

    _add_heading(doc, "2. Feature Engineering", 1)
    _add_heading(doc, "2.1 Pre-Start Proactivity (Negative VLE Dates)", 2)
    _add_body(
        doc,
        "Negative dates count days before the official module start. We aggregate pre_start_clicks, "
        "pre_start_active_days, days_before_start, and material-study signals into p_pre_start_proactivity."
    )
    for name in ["pre_start_clicks", "days_before_start", "pre_start_active_days", "pre_start_proactivity_raw"]:
        if name in feature_rationale["feature"].values:
            _add_body(doc, feature_description(name, "", feature_rationale))
    _add_body(
        doc,
        f"Population check: {pre_start['with_pre_start']:,} of {pre_start['enrolments']:,} enrolments "
        f"({pre_start['with_pre_start_pct']:.1%}) show pre-start activity. Risk rate with pre-start: "
        f"{pre_start['risk_with_pre_start']:.1%}; without: {pre_start['risk_without_pre_start']:.1%}."
    )
    _add_body(
        doc,
        "By contrast, date_unregistration is treated as an audit field rather than a predictive feature. It can reveal "
        "future withdrawal timing, so including it in the Week 6 model would overstate real-world performance."
    )

    _add_heading(doc, "2.2 In-Semester VLE and Assessment Features", 2)
    _add_body(
        doc,
        "Task 1 features are grouped into six interpretable buckets: activity volume, consistency and rhythm, "
        "recency and disengagement gaps, resource diversity and intent, course-pace alignment, and assessment "
        "behaviour and recovery."
    )
    for name, desc in [
        ("weekly_clicks_norm", "Log-scaled click volume, peer-normalized within module/presentation/week."),
        ("unique_sites_wk", "Distinct VLE sites accessed in the week."),
        ("active_days_last_7", "Distinct active days in the week."),
        ("study_regularity_score", "Active days divided by seven."),
        ("burstiness_score", "Share of weekly clicks concentrated on the busiest day."),
        ("week_to_week_volatility", "Rolling instability in weekly click volume."),
        ("days_since_last_click", "Recency of last observed click."),
        ("inactive_last_14_days", "Flag for two-week absence."),
        ("longest_inactive_gap_so_far", "Largest inactivity gap observed so far."),
        ("activity_diversity", "Distinct VLE activity types used."),
        ("activity_entropy_wk", "Balance of clicks across VLE activity types."),
        ("material_active_days", "Frequency of learning-material study activity."),
        ("material_click_share", "Share of clicks focused on learning materials."),
        ("recent_material_active_days_2w", "Recent learning-material study rhythm near Week 6."),
        ("forum_active_days", "Discussion-board/help-seeking participation."),
        ("on_schedule_click_ratio", "Metadata-covered clicks that align with the planned course weeks."),
        ("planned_material_coverage", "Cumulative share of expected planned materials accessed."),
        ("assessment_submitted_ratio", "Cumulative non-exam assessment completion."),
        ("missed_assessments_cum", "Cumulative assessments due but not submitted."),
        ("pre_assessment_clicks_7d_cum", "Preparation clicks in the seven days before due dates."),
        ("cram_ratio_week", "Share of preparation clicks occurring in the final two days."),
        ("low_weight_completion_ratio", "Low-weight assignment completion."),
    ]:
        if name in feature_rationale["feature"].values:
            _add_body(doc, feature_description(name, desc, feature_rationale))
    _add_figure(doc, FIG_DIR / "week6_feature_rationale.png", "Figure 3. Week 6 feature correlations with risk.")

    _add_heading(doc, "3. Dynamic Engagement Score", 1)
    _add_body(
        doc,
        "Ten peer-normalized components (including pre-start proactivity) are combined with train-derived weights "
        "into a 0–100 weekly score."
    )
    weight_rows = [
        [row["label"], f"{row['success_auc']:.3f}", f"{100 * row['score_weight']:.0f}%"]
        for _, row in weights_df.iterrows()
    ]
    _add_table(doc, ["Component", "Success AUC", "Final weight"], weight_rows)
    _add_body(
        doc,
        "Scorecard guardrail: any component with train success AUC at or below 0.50 receives 0% weight. "
        "Those signals can still be inspected as exploratory features, but they do not move the trusted "
        "administrator-facing engagement score."
    )

    band_rows = [
        [row["score_band"], int(row["enrolments"]), f"{row['observed_withdraw_fail_rate']:.1%}", f"{row['avg_score']:.1f}"]
        for _, row in score_bands.iterrows()
    ]
    _add_table(doc, ["Score band", "Enrolments", "Withdraw/fail rate", "Avg score"], band_rows)
    _add_body(
        doc,
        f"Week 6 engagement score correlation with risk: r = {engagement_row['risk_correlation']:+.3f}."
    )
    _add_figure(doc, FIG_DIR / "score_band_risk.png", "Figure 4. Risk by engagement score band.")
    archetype_def_rows = [
        [row["archetype"], row["description"], row["feature_signature"]]
        for _, row in archetype_definitions().iterrows()
    ]
    _add_table(doc, ["Archetype", "Description", "Feature signature"], archetype_def_rows)
    _add_body(doc, "Representative archetypes: " + "; ".join(archetype_lines()) + ".")
    _add_figure(doc, FIG_DIR / "engagement_archetypes_core.png", "Figure 5a. Core engagement archetypes.")
    _add_figure(doc, FIG_DIR / "engagement_archetypes_additional.png", "Figure 5b. Additional behavioral archetypes.")

    _add_heading(doc, "4. Week 6 Disengagement Model", 1)
    _add_body(
        doc,
        f"Binary classifier: Withdrawn/Fail vs Pass/Distinction. Stratified holdout: "
        f"{split['test']:,} test enrolments ({split['test'] / split['total']:.1%} of labelled Week 6 cohort)."
    )
    _add_body(
        doc,
        "The notebook model-selection table compared Logistic Regression, Random Forest, XGBoost, and a "
        "regularized XGBoost variant. XGBoost is now the production model because it delivered the best "
        "validation F1 and recall while keeping ROC-AUC and PR-AUC competitive. Thresholding is a separate "
        "product decision: after selecting the model, operating cutoffs are chosen from the precision-recall "
        "tradeoff based on advisor workload. To reduce runtime, default pipeline runs now train XGBoost only "
        "and retain the comparison table as model-selection evidence."
    )
    comparison_rows = [
        [
            row["model_name"],
            f"{row['precision']:.3f}",
            f"{row['recall']:.3f}",
            f"{row['f1']:.3f}",
            f"{row['roc_auc']:.3f}",
            row.get("selection_note", ""),
        ]
        for _, row in model_comparison.head(4).iterrows()
    ]
    _add_table(doc, ["Model", "Precision", "Recall", "F1", "ROC-AUC", "Decision"], comparison_rows)
    _add_figure(
        doc,
        FIG_DIR / "task2_behavioral_correlation_heatmap.png",
        "Figure 6. Task 2 Week 6 behavioural feature correlation heatmap.",
        width=6.4,
    )
    _add_body(
        doc,
        "Task 2 now carries forward the Task 1 engagement framework directly: the model includes engagement_score_6, "
        "a Week 6-safe behavioural archetype, train-only archetype/profile risk priors, ordered education and age "
        "features, IMD midpoint, module-normalized credit load, resource-mix ratios, and selected interaction terms."
    )
    thresh_rows = [
        [
            row["threshold_name"],
            f"{row['threshold']:.3f}",
            f"{row['precision']:.3f}",
            f"{row['recall']:.3f}",
            int(row["false_negatives"]),
            int(row["false_positives"]),
        ]
        for _, row in threshold_table.iterrows()
    ]
    _add_table(doc, ["Threshold", "Value", "Precision", "Recall", "FN", "FP"], thresh_rows)
    _add_body(
        doc,
        f"High-touch queue (threshold {urgent['threshold']:.3f}): precision {urgent['precision']:.3f}, "
        f"recall {urgent['recall']:.3f}, students {urgent['alerts']} ({urgent['alert_rate']:.1%} of test cohort). "
        f"The broader top-60% light-touch tier has F1 {watchlist['f1']:.3f} and recall {watchlist['recall']:.3f}; "
        f"overall model ranking quality is ROC-AUC {urgent['roc_auc']:.3f}. The model ranks risk; the product decides "
        "whether that risk becomes high-touch support, light-touch nudges, or monitoring."
    )
    _add_heading(doc, "4.1 What An Advisor Receives", 2)
    _add_body(
        doc,
        "At the end of Week 6, each course presentation generates a ranked list of students needing support. For each "
        "student, the advisor sees the student ID, course, predicted risk score, intervention tier, three plain-English "
        "reason codes, and a recommended action."
    )
    _add_body(
        doc,
        "Reason codes are translated from the strongest per-student risk signals, for example: missed or late "
        "assessments, silent weeks in the six-week window, a long gap since last VLE activity, low activity diversity, "
        "or low assessment completion. The advisor sees why this student is flagged, not just a probability."
    )
    _add_body(
        doc,
        "Delivery is tiered: high-touch students appear in a live advisor queue and trigger a push notification; "
        "light-touch students are included in a weekly digest and automated nudge workflow; monitoring students remain "
        "visible on the dashboard without immediate action."
    )
    tier_rows = [
        [
            row["tier"],
            int(row["students"]),
            f"{row['student_share']:.1%}",
            f"{row['observed_withdraw_fail_rate']:.1%}",
            f"{row['captured_risk_share']:.1%}",
        ]
        for _, row in intervention_tiers.iterrows()
    ]
    _add_table(doc, ["Tier", "Students", "Share", "Observed risk", "Captured risk"], tier_rows)
    _add_figure(doc, FIG_DIR / "intervention_tiers.png", "Figure 7. Capacity-based intervention tiers.")
    _add_figure(doc, FIG_DIR / "confusion_matrix.png", "Figure 8. Top-60% light-touch confusion matrix.")
    _add_figure(doc, FIG_DIR / "threshold_tradeoff.png", "Figure 9. Risk cutoff tradeoff.")
    _add_figure(doc, FIG_DIR / "calibration.png", "Figure 10. Calibration curve.")

    mechanisms = {
        "avg_score_so_far_6": "Early academic struggle may signal content difficulty before final failure.",
        "engagement_score_6": "Combines recency, consistency, diversity, and assessment behavior into a trajectory signal.",
        "assessment_submitted_ratio_6": "Missing early assessments is both predictive and directly actionable for advisors.",
    }
    overall_drivers = _load_csv("risk_feature_drivers.csv").head(3)
    driver_rows = [
        ["Overall", row["feature"].replace("num__", ""), mechanisms.get(row["feature"].replace("num__", ""), "High-ranking model driver.")]
        for _, row in overall_drivers.iterrows()
    ]
    driver_rows.extend(
        [
            [
                "Behavioral",
                row["feature"].replace("num__", ""),
                mechanisms.get(row["feature"].replace("num__", ""), "Advisor-actionable engagement or assessment behavior."),
            ]
            for _, row in behavioral_drivers.head(3).iterrows()
        ]
    )
    _add_table(doc, ["Driver group", "Feature", "Mechanism"], driver_rows)
    _add_body(
        doc,
        f"Top driver: assessment completion (withdraw/fail mean {assess_row['withdraw_fail_mean']:.1%} vs "
        f"{assess_row['pass_distinction_mean']:.1%} for successful students). "
        "The advisor-facing output includes tier, calibrated probability, engagement trajectory, assessment context, "
        "and top drivers."
    )
    _add_table(
        doc,
        ["PathAI Advisor Alert", "Example"],
        [
            ["Student", "S-10482"],
            ["Course", "DDD-2014J"],
            ["Risk tier", "High-touch support queue"],
            ["Predicted risk", "78%"],
            [
                "Why flagged",
                "No VLE activity in 12 days; engagement score dropped from 61 to 34 over two weeks; first assessment due by Week 6 was not submitted.",
            ],
            [
                "Suggested action",
                "Prioritize for advisor outreach. Ask about workload, access issues, and confidence with course material. Offer academic support session.",
            ],
        ],
    )

    _add_heading(doc, "5. Course Recommendation Engine", 1)
    _add_body(
        doc,
        "Primary user: students planning next semester. Content-based scoring uses demographics, engagement band, "
        "and pre-start proactivity history. Collaborative filtering uses cosine similarity with a boost for proactive peers."
    )
    rec_rows = [
        ["Holdout students", rec_metrics["holdout_students"]],
        ["Content hit@3", f"{rec_metrics['content_hit_rate_at_3']:.1%}"],
        ["CF hit@3", f"{rec_metrics['cf_hit_rate_at_3']:.1%}"],
        ["Cold start (proactive)", ", ".join(rec_metrics["cold_start_strategy"])],
    ]
    _add_table(doc, ["Metric", "Value"], rec_rows)

    _add_heading(doc, "6. Reproducibility", 1)
    _add_body(
        doc,
        "Run: python3 -m venv .venv && .venv/bin/pip install -r requirements.txt && "
        ".venv/bin/python src/run_pipeline.py. CSV/JSON artifacts go to outputs/; figures and reports to reports/."
    )

    doc.save(output_path)
    return output_path


def main() -> None:
    path = write_word_report()
    print(f"Word report written to: {path}")


if __name__ == "__main__":
    main()
